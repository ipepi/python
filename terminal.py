import os
my_files = ['accounts.txt', 'detail.csv', 'invite.docx']
for filenames in my_files:
    print(os.path.join('usr/bin/spam',filenames))


import shelve
shelf_file = shelve.open('mydata')
cats = ['Zophie', 'Pooka', 'Shimon']
shelf_file['cats'] = cats
shelf_file.close



shelf_file = shelve.open('mydata')
type(shelf_file)
shelf_file['cats']
shelf_file.close


import pprint
cats = [{'name': 'Zophie', 'desc': 'chubby'}, {'name': 'Pooka', 'desc': 'fluffy'}]
pprint.pformat(cats)

file_obj = open('myCats.py', 'w')
file_obj.write('cats = ' + pprint.pformat(cats) + '\n')
file_obj.close

import myCats
myCats.cats

import os
for foldername, subfolders, filenames in os.walk('/Users/Ippei/python'):
    print('The current folder is ' + foldername)

    for subfolder in subfolders:
        print('SUBFOLDER OF ' + foldername + ': ' + subfolder)
    for filename in filenames:
        print('FILE INSIDE ' + foldername + ': ' + filename)
        print(' ')


import zipfile, os
os.chdir('/Users/Ippei/python/folder')
example_zip = zipfile.ZipFile('example.zip')

import zipfile, os
os.chdir('/Users/Ippei/python/folder')
example_zip = zipfile.ZipFile('example.zip')
example_zip.namelist()

spam_info = example_zip.getinfo('spam.txt')
spam_info.file_size

example_zip.extractall()

import zipfile
new_zip = zipfile.ZipFile('new.zip', 'w')
new_zip.write('spam.txt', compress_type=zipfile.ZIP_DEFLATED)
new_zip.close()

import requests, bs4
res = requests.get('http://nostarch.com')
res.raise_for_status()
no_starch_soup = bs4.BeautifulSoup(res.text)
type(no_starch_soup)

import requests, bs4
example_file = open('example.html')
example_soup = bs4.BeautifulSoup(example_file)
type(example_soup)

elems = example_soup.select('#author')
type(elems)
len(elems)
type(elems[0])
elems[0].getText()
str(elems[0])
elems[0].attrs


from selenium import webdriver
browser = webdriver.Firefox()
type(browser)
browser.get('http://inventwithpython.com')
link_elem = browser.find_element_by_link_text('Read Online for Free')
type(link_elem)
link_elem.click()

from seleinium import webdriver
from selenium.webdriver.common.keys import Keys
browser = webdriver.Firefox()
browser.get('http://nostarch.com')
html_elem = browser.find_element_by_class_name('New game')
html_elem.send_keys(Keys.END)
html_elem.send_keys(Keys.HOME)


import openpyxl
wb = openpyxl.load_workbook('exampe.xlsx')
wb.get_sheet_names()
sheet = wb.get_sheet_by_name('Sheet1')
sheet.cell(row=1, column=2)

for i in range(1, 8, 2):
    print(i, sheet.cell(row=i, column=2).value)



import openpyxl
wb = openpyxl.load_workbook('example.xlsx')
sheet = wb.get_sheet_by_name('Sheet1')
tupple(sheet['A1':'C3'])


import openpyxl
wb = openpyxl.Workbook()
wb.get_sheet_names()
sheet = wb.active
sheet.title
sheet.title = 'Spam Bacon Rggs Sheet'
wb.get_sheet_names()


import openpyxl
wb = openpyxl.load_workbook('example.xlsx')
sheet = wb.active
sheet.title = 'Spam Spam Spam'
wb.save('example_copy.xlsx')

import openpyxl
wb = openpyxl.Workbook()
wb.get_sheet_names()
wb.create_sheet()
wb.get_sheet_names()
wb.create_sheet(index=0, title='First Sheet')
wb.get_sheet_names()
wb.create_sheet(index=2, title='Middle Sheet')
wb.get_sheet_names()
wb.remove_sheet(wb.get_sheet_by_name('Middle Sheet'))
wb.remove_sheet(wb.get_sheet_by_name('Sheet1'))
wb.get_sheet_names()

import openpyxlwb = openpyxl.Workbook()
wb = openpyxl.Workbook()
sheet = wb.get_sheet_by_name('Sheet')
sheet['A1'] = 'Hello World'
sheet['A1'].value


import openpyxl
wb = openpyxl.Workbook()
sheet = wb.active
sheet['A1'] = 200
sheet['A2'] = 300
sheet['A3'] = '=SUM(A1:A2)'
wb.save('writeFormula.xlsx')

import openpyxl
wb = openpyxl.Workbook()
sheet = wb.active
sheet.merge_cells('A1:D3')
sheet['A1'] = 'Tweleve cells merged together'
sheet.merge_cells('C5:D5')
sheet['C5'] = 'Two merged cells.'
wb.save('merged.xlsx')

import openpyxl
wb = openpyxl.load_workbook('merged.xlsx')
sheet = wb.active
sheet.unmerge_cells('A1:D3')
sheet.unmerge_cells('C5:D5')
wb.save('merged.xlsx')

import openpyxl
wb = openpyxl.load_workbook('produceSales.xlsx')
sheet = wb.active
sheet.freeze_panes = 'A2'
wb.save('freezeExample.xlsx')

import openpyxl
wb = openpyxl.Workbook()
sheet = wb.active
for i in range(1,11):
    sheet['A' + str(i)] = i

ref_obj = openpyxl.chart.Reference(sheet, min_col=1, min_row=1, max_col=1, max_row=10)
series_obj = openpyxl.chart.Series(ref_obj, title='First series')
chart_obj = openpyxl.chart.BarChart()
chart_obj.append(series_obj)
#chart_obj.y = 10
#chart_obj.x = 100
#chart_obj.w = 300
#chart_obj.h = 200
sheet.add_chart(chart_obj)
wb.save('sampleChart.xlsx')


import PyPDF2
pdf1_file = open('meetingminutes.pdf', 'rb')
pdf2_file = open('meetingminutes2.pdf', 'rb')
pdf1_reader = PyPDF2.PdfFileReader(pdf1_file)
pdf2_reader = PyPDF2.PdfFileReader(pdf2_file)
pdf_writer = PyPDF2.PdfFileWriter()
for page_num in range(pdf1_reader.numPages):
    page_obj = pdf1_reader.getPage(page_num)
    pdf_writer.addPage(page_obj)


for page_num in range(pdf2_reader.numPages):
    page_obj = pdf2_reader.getPage(page_num)
    pdf_writer.addPage(page_obj)


pdf_output_file = open('combinedminutes.pdf', 'wb')
pdf_writer.write(pdf_output_file)
pdf_output_file.close()
pdf1_file.close()
pdf2_file.close()


import PyPDF2
minutes_file = open('meetingminutes.pdf', 'rb')
pdf_reader = PyPDF2.PdfFileReader(minutes_file)
page = pdf_reader.getPage(0)
page.rotateClockwise(90)

pdf_writer = PyPDF2.PdfFileWriter()
pdf_writer.addPage(page)
result_pdf_file = open('rotatedPage.pdf', 'wb')
pdf_writer.write(result_pdf_file)
result_pdf_file.close()
minutes_file.close()



import PyPDF2
minutes_file = open('meetingminutes.pdf', 'rb')
pdf_reader = PyPDF2.PdfFileReader(minutes_file)
minutes_first_page = pdf_reader.getPage(0)
pdf_watermark_reader = PyPDF2.PdfFileReader(open('watermark.pdf', 'rb'))
minutes_first_page.mergePage(pdf_watermark_reader.getPage(0))
pdf_writer = PyPDF2.PdfFileWriter()
pdf_writer.addPage(minutes_first_page)
for page_num in range(1, pdf_reader.numPages):
    page_obj = pdf_reader.getPage(page_num)
    pdf_writer.addPage(page_obj)

result_pdf_file = open('watermarkedCover.pdf', 'wb')
pdf_writer.write(result_pdf_file)
minutes_file.close()
result_pdf_file.close()




import PyPDF2
minutes_file = open('meetingminutes.pdf', 'rb')
pdf_reader = PyPDF2.PdfFileReader(minutes_file)
pdf_writer = PyPDF2.PdfFileWriter()
for page_num in range(pdf_reader.numPages):
    pdf_writer.addPage(pdf_reader.getPage(page_num))

pdf_writer.encrypt('swordfish')
result_pdf = open('encryptedminutes.pdf', 'wb')
pdf_writer.write(result_pdf)
result_pdf.close()


import docx
doc = docx.Document('demo.docx')
len(doc.paragraphs)
doc.paragraphs[0].text
doc.paragraphs[1].text
len(doc.paragraphs[1].runs)

doc.paragraphs[1].runs[0].text
doc.paragraphs[1].runs[1].text
doc.paragraphs[1].runs[2].text
doc.paragraphs[1].runs[3].text


import csv
example_file = open('example.csv')
example_reader = csv.reader(example_file)
example_data = list(example_reader)
example_data


import csv
example_file = open('example.csv')
example_reader = csv.reader(example_file)
for row in example_reader:
    print('ROW #' + str(example_reader.line_num) + ' ' + str(row))



import csv
output_file = open('output.csv', 'w', newline='')
output_writer = csv.writer(output_file)
output_writer.writerow(['spam', 'egg', 'bacon', 'ham'])
output_writer.writerow(['Hello, world!', 'eggs', 'bacon', 'ham'])
output_writer.writerow([1, 2, 3.131592, 4])
output_file.close()

import csv 
csv_file = open('example.tsv', 'w', newline='')
csv_writer = csv.writer(csv_file, delimiter='\t', lineterminator='\n\n')
csv_writer.writerow(['apples', 'oranges', 'grapes'])
csv_writer.writerow(['eggs', 'bacon', 'ham'])
csv_writer.writerow(['spam', 'spam', 'spam', 'spam', 'spam', 'spam'])
csv_file.close()


string_of_json_data = '{"name": "Zophie", "isCat": true, "miceCaught": 0, "felineIQ": null}'
import json
json_data_as_python_value = json.loads(string_of_json_data)
json_data_as_python_value

import datetime
import time
halloween2020 = datetime.datetime(2020, 10, 31, 0, 0, 0)
while datetime.datetime.now() < halloween2020:
    time.sleep(1)

oct21st = datetime.datetime(2015, 10, 21, 16, 29, 0)
oct21st.strftime('%Y/%m/%d %H:%M:%S')

import threading, time
print('プログラムの開始')

def take_a_nap():
    time.sleep(5)
    print('起きた！')

thread_obj = threading.Thread(Target=take_a_nap)
thread_obj.start()

print('プログラム終了')